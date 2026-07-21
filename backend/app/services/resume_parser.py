"""Safe in-memory PDF and DOCX resume text extraction."""

import re
import xml.etree.ElementTree as element_tree
import zipfile
from io import BytesIO

import fitz
from docx import Document

from app.schemas.resume import ResumeFileType, ResumeUploadResponse
from app.utils.file_validation import ValidatedResumeUpload


class ResumeParseError(ValueError):
    """Raised when a supported file cannot be parsed as its declared document type."""


def _normalise_text(text: str) -> str:
    """Remove noisy whitespace while preserving readable paragraph boundaries."""
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def _word_count(text: str) -> int:
    """Return a Unicode-aware count of words in extracted text."""
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))


def _parse_pdf(content: bytes) -> tuple[str, int]:
    """Extract text and page count from a PDF document using PyMuPDF."""
    try:
        document = fitz.open(stream=content, filetype="pdf")
    except (fitz.FileDataError, RuntimeError) as error:
        raise ResumeParseError("The uploaded PDF could not be opened.") from error

    try:
        if document.page_count == 0:
            raise ResumeParseError("The uploaded PDF has no pages.")
        return "\n\n".join(page.get_text("text") for page in document), document.page_count
    finally:
        document.close()


def _docx_page_count(content: bytes) -> int | None:
    """Read Word's saved page count from extended document properties when available."""
    namespace = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}"
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            properties = element_tree.fromstring(archive.read("docProps/app.xml"))
        page_element = properties.find(f"{namespace}Pages")
        if page_element is None or not page_element.text:
            return None
        page_count = int(page_element.text)
        return page_count if page_count > 0 else None
    except (KeyError, ValueError, zipfile.BadZipFile, element_tree.ParseError):
        return None


def _parse_docx(content: bytes) -> tuple[str, int | None]:
    """Extract paragraph and table-cell text from a DOCX document using python-docx."""
    try:
        document = Document(BytesIO(content))
    except (ValueError, zipfile.BadZipFile, KeyError) as error:
        raise ResumeParseError("The uploaded DOCX could not be opened.") from error

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join([*paragraphs, *table_cells]), _docx_page_count(content)


def parse_resume(upload: ValidatedResumeUpload) -> ResumeUploadResponse:
    """Parse validated resume bytes and return normalized text with document metadata."""
    try:
        if upload.file_type is ResumeFileType.PDF:
            extracted_text, page_count = _parse_pdf(upload.content)
        elif upload.file_type is ResumeFileType.DOCX:
            extracted_text, page_count = _parse_docx(upload.content)
        else:
            raise ResumeParseError("The uploaded file type is not supported.")
    except ResumeParseError:
        raise
    except Exception as error:
        raise ResumeParseError("The uploaded resume could not be parsed.") from error

    raw_text = _normalise_text(extracted_text)
    if not raw_text:
        raise ResumeParseError("No readable text could be extracted from the uploaded resume.")

    return ResumeUploadResponse(
        filename=upload.filename,
        file_type=upload.file_type,
        file_size_bytes=upload.file_size_bytes,
        page_count=page_count,
        character_count=len(raw_text),
        word_count=_word_count(raw_text),
        raw_text=raw_text,
    )
